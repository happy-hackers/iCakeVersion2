# -*- coding: UTF-8 -*-  
import sys     
reload(sys)
sys.setdefaultencoding('utf-8')

try:
    from tkinter import *
except ImportError:
    import Tkinter
import webbrowser

try:
    from urllib import quote  # Python 2.X
except ImportError:
    from urllib.parse import quote  # Python 3+
import ttk
from multicolumnlistboxLib import *
from database import *
#from ttkcalendar import *

from shutil import copyfile

from datetime import datetime
import googlemaps
#from kmeans import *
from PIL import Image, ImageTk
from autocomplete import AutocompleteEntry

#import shopify
#import requests
import re

API_KEY = 'ffc3a9728b3e4f62be551fe88c99e83e'
PASSWORD = 'e269e4af0f1a66f9248f576e09db35ee'
SHOP_NAME = 'icake-melbourne'

shop_url = "https://%s:%s@%s.myshopify.com/admin" % (API_KEY, PASSWORD, SHOP_NAME)


lista = ['36 latrob st', 'actions', 'additional', 'also', 'an', 'and', 'angle', 'are', 'as', 'be', 'bind', 'bracket', 'brackets', 'button', 'can', 'cases', 'configure', 'course', 'detail', 'enter', 'event', 'events', 'example', 'field', 'fields', 'for', 'give', 'important', 'in', 'information', 'is', 'it', 'just', 'key', 'keyboard', 'kind', 'leave', 'left', 'like', 'manager', 'many', 'match', 'modifier', 'most', 'of', 'or', 'others', 'out', 'part', 'simplify', 'space', 'specifier', 'specifies', 'string;', 'that', 'the', 'there', 'to', 'type', 'unless', 'use', 'used', 'user', 'various', 'ways', 'we', 'window', 'wish', 'you']

# gmaps = googlemaps.Client(key="AIzaSyB3Ao6QZnqxngkZPB4d5yQaboPp-mSjf4s")
key = ["AIzaSyCWzV0pbt_84I2DraGqg1OaC5kil5pZESY",
        "AIzaSyD44UyOsGzuyngSnb2WPLPuFGzhIG1OL1s",
        "AIzaSyB3Ao6QZnqxngkZPB4d5yQaboPp-mSjf4s",
        "AIzaSyBykVZQJbt518Jh58CDo6vb3TH4pM0j21Q"]

prefix = "https://www.google.com/maps/dir/?api=1"
origin = "&origin="
destination = "&destination="
viaDriving = "&travelmode=driving"
waypoints = "&waypoints="

num_key = 0
gmaps = googlemaps.Client(key=key[num_key])
window_open = False                  # to ensure only one add/edit order window
                                     # can be opened at one time
cake_window_open = False             # same functionality as window_open
popup_cake = None                    # right click menu for cakes windows

# calculate distance between two points using google map
def cal_dis(orders,end_loc):
    if orders:
            locations = to_loc(orders)
            locations.insert(0,proc_start_loc.get())

            # if there is specified end location append it to the end of locations
            write_2_log("cal_dis, working...")
            write_2_log("end location: {}".format(end_loc))
            write_2_log(','.join(get_address(end_loc)))
            
            
            print "end location: {}".format(end_loc)            
            print ','.join(get_address(end_loc))
            if end_loc != hint1 and ','.join(get_address(end_loc)):
                print "$"                
                locations.append(','.join(get_address(end_loc)))
                print "locations"
                print locations
                write_2_log("locations")
                write_2_log(locations)
                
                
                
            directions_result = gmaps.directions(locations[0],locations[-1],
                                                waypoints=locations[1:-1],
                                                mode="driving", avoid="tolls",
                                                departure_time=datetime.now(),
                                                optimize_waypoints = True)
            totalDistance = 0
            totalDuration = 0
            newLocationIds = []
            newLocations = {}
            
            if not directions_result:
                write_2_log("     请检查地址，一些地址的内容有错!     ")
                warning_window(master,"     请检查地址，一些地址的内容有错!     ")
                return
            lists = directions_result[0]['legs']
            # loop through the details of direction results
            write_2_log("range(len(lists) = " + str(len(lists)))
            write_2_log(lists[0]['start_address'])
            write_2_log(lists[0]['end_address'])
            # write_2_log(lists[1]['end_address'])
#             write_2_log(lists[2]['end_address'])
                
            for index in range(len(lists)):
                write_2_log("index = {}".format(index))
                
                try:
                    #print lists[index]
                    if lists[index]['distance']['text'][:-3]:
                        totalDistance += float(lists[index]['distance']['text'][:-3])
                    if lists[index]['duration']['text'][:-4]:
                        totalDuration += float(lists[index]['duration']['text'][:-4])
                except IndexError as e:
                    write_2_log(e)
                    print e
                    return
                
                if index == 0:
                    locationId = gmaps.geocode(lists[index]['start_address'])[0]['place_id']
                    newLocations[locationId] = lists[index]['start_address']
                    newLocationIds.append(locationId)
                     
                write_2_log("line = {}".format(133))   
                locationId = gmaps.geocode(lists[index]['end_address'])[0]['place_id']
                write_2_log("line = {}".format(134))
                newLocations[locationId] = lists[index]['end_address']
                write_2_log("line = {}".format(135))    
                write_2_log("\n######location: {}\nid:{}".format(lists[index]['end_address'],locationId))
                write_2_log(lists[index]['end_address'])
                newLocationIds.append(locationId)
                
            write_2_log(newLocationIds)
            print newLocationIds
            finalLocations = outputFile(newLocationIds, totalDistance, \
                                     totalDuration, (len(locations)-1), orders,end_loc,newLocations)
            openWebsite(finalLocations)    

    else:
        return
        
# find order by id helper
def find_order_by_id2(orders,id):
    for order in orders:
        if str(order.order_number) == str(id):
            return order

# setup doc
def setup_doc(document,order,order_num):
    strr = u"订单号:{},{}".format(order.order_number,
                                           order.phone)
    if order.name:
        strr = strr + " {}收".format(order.name)
    if order.price:
        strr = strr + " ${}".format(order.price)
    if order.upfront:
        strr = strr + "(${} paid)".format(order.upfront)
    if order.writing:
        strr = strr + " {}".format(order.writing)
    if order.candle:
        strr = strr + " {}".format(order.candle)
    if order.tableware:
        strr = strr + " {}".format(order.tableware)
    if order.ps_info:
        strr = strr + " {}".format(order.ps_info)
               
    document.add_paragraph(strr)
    for cake in order.cake_type:
        document.add_paragraph(u"{}) {} {} {} {}".format(cake[index_cake_no],
                                                     cake[index_cake_type],
                                                     cake[index_cake_size],
                                                     cake[index_cake_inner],
                                                     cake[index_cake_ps]))
    
# setup doc table
def setup_table(document,order):
    cake_num = 1
    row_num = 0
    table = document.add_table(rows = (len(order.cake_type)+1), cols = 5)
    table.style = 'Table Grid'

    # setup header for the table
    row = table.rows[row_num]
    row.cells[0].text = u"No"
    row.cells[1].text = u"蛋糕款式"
    row.cells[2].text = u"蛋糕尺寸"
    row.cells[3].text = u"蛋糕内芯"
    row.cells[4].text = u"蛋糕备注"

    row_num += 1
    for cake in order.cake_type:
        row = table.rows[row_num]
        row.cells[0].text = str(cake_num)

        # bold cell
        cell_type = row.cells[1]
        cell_type.text = cake[index_cake_type]
        run = cell_type.paragraphs[0].runs[0]
        run.font.bold = True
        # bold cell
        cell_size = row.cells[2]
        cell_size.text = cake[index_cake_size]
        run = cell_size.paragraphs[0].runs[0]
        run.font.bold = True

        row.cells[3].text = cake[index_cake_inner]
        row.cells[4].text = cake[index_cake_ps]
        
        row_num += 1
        cake_num += 1
               
###############################################################################
def outputFile(finalLocationIds, totalDistance, totalDuration, totalOrder, \
        orders,end_loc,newLocations):
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    import easygui
    import os

    document = Document()

    dic_id = {}
    # input order info
    if end_loc != hint1:
        # driver has address
        if ','.join(get_address(end_loc)):
            print "#####"
            print "end_loc is"
            print end_loc
            end_loc_without_name = (','.join(get_address(end_loc)))
            print end_loc_without_name
            document.add_paragraph("Dispatcher:{}".format(end_loc,font=("Calibri",title_size)))
            end_loc_id = gmaps.geocode(gmaps.places_autocomplete(\
                    input_text = end_loc_without_name, \
                        location = end_loc_without_name)[0]['description'])[0]['place_id']

            dic_id[end_loc_id] = "end location"
        else:
            document.add_paragraph("Dispatcher:{}".format(end_loc.split(',')[0],font=("Calibri",title_size)))  
    else:
        document.add_paragraph("Driver Not Determined".format(font=("Calibri",title_size)))

    finalLocations = []

    dic_id[finalLocationIds[0]] = "start location"
    for order in orders:
        id = gmaps.geocode(order.address)[0]['place_id']
        try:
            dic_id[id] = str(dic_id[id]) + "," + str(order.order_number)
        except KeyError:
            dic_id[id] = str(order.order_number)      
        print "dic_id[{}] = {}".format(id,dic_id[id])
        
    print dic_id
    print "ouput:"
    print finalLocationIds
    print unique(finalLocationIds)
    order_num = 0
    order_num2 = 0
    for item in unique(finalLocationIds):
        finalLocations.append(newLocations[item])
        print item
        try:
            if order_num != 0:
                order_numberr = dic_id[item]
                print order_numberr
                document.add_paragraph("{} : {}".format(order_num,newLocations[item], level = 1))
                order_num += 1
                for num in order_numberr.split(","):
                    order = find_order_by_id2(orders,num)
                    print "num = {},order:".format(num)
                    print order
                    if order:
                        order_num2 += 1
                        setup_doc(document,order,order_num)
            else:
               order_num += 1 
                
             
        except KeyError:
            warning_window(master,"{}\n地址存在问题.".format(newLocations[item]))
        
       
            
    #document.add_paragraph("\nTotal distance is %.1f kms." % totalDistance)
    #document.add_paragraph("Total duration is %.1f mins." % totalDuration)      #document.add_paragraph(u"\n总共{}个订单，{}个需要送货的地址.\n".format(order_num2,(order_num-1)))
    
    write_2_log("outputfile, working...")
    
    now = datetime.now()
    path = easygui.diropenbox()

    if end_loc != hint1:
        filename = now.strftime("%Y-%m-%d %H.%M ")+"for {}".format(end_loc.split(',')[0])+".docx"
    else:
        filename = now.strftime("%Y-%m-%d %H.%M ")+"for {}".format("Driver")+".docx"

    filepath = os.path.join(path, filename)
    document.save(filepath)

    return finalLocations

###############################################################################
def openWebsite(finalLocations):
    i = 0
    print "#####################"
    print finalLocations
    addPage = False
    while i < len(finalLocations):
        if (i % 10 == 0):
            if i != 0:
                webbrowser.open(prefix+origin+ start_point +destination+ end_point +\
                        viaDriving + waypoints + quote(points[:-1]) )
            points = ""
            start_point = quote(finalLocations[i])

        # if i is every 9th index item or the last one for all locations
        elif ((i + 1) % 10 == 0 or i == len(finalLocations)-1):
            end_point = quote(finalLocations[i])
            if (i == len(finalLocations)-1):
                webbrowser.open(prefix+origin+ start_point +destination+ end_point +\
                        viaDriving + waypoints + quote(points[:-1]) )

        else:
            #There will be one more | at the end
            points += finalLocations[i]+"|"

        i += 1

###############################################################################
def find_order_by_id(id,orders):
    for order in orders:
        if int(id) == int(order.order_number):
            return order
    print "Error(find_order_by_id()): no order found"
    return None

def to_loc(orders):
    locs = []
    for i in orders:
        locs.append(i.address)
    return locs

def to_order(orders):
    order_numbers = []
    for i in orders:
        order_numbers.append(i.order_number)
    return order_numbers

def to_order2(orders):
    order_numbers = []
    for i in orders:
        order_numbers.append([i.order_number])
    return order_numbers


def get_item_by_num(num,orders):
    for i in orders:
        new = []
        for n in  num.split(',')[0]:
            if n.isdigit() or n.isalpha():
                new.append(n)
        strr = ''.join(new)
        print strr
        print i.order_number
        if str(i.order_number) == str(strr):
            return i

def get_item_by_num2(num,orders):
    print "get_item_by_num2"
    if not orders:
        print "orders is null"
    for i in orders:
        print i
        print "i.order_number = {}".format(i.order_number)
        print "num[0] = {}".format(num)
        if int(i.order_number) == int(num):
            return i
###############################################################################
# check whether its valid address or not
def isvalid_ad(address):
    count = 0
    for i in address:
        if i == ',':
            count += 1
        if i == '，':
            return 2
    if count != 1:
        return -1
    # postcode is not in the database
    else:        
        f = in_pos(address.split(",")[1],db.postcode[index_city],db.postcode_full[index_city])
        s = in_pos(address.split(",")[1],db.postcode[index_other],db.postcode_full[index_other])
        if (not f) and (not s):
            return 1      
    return 0
    
# sperate orders to orders in city and orders in other areas
def in_city(orders,postcode_city,postcode_full_city):
    return_list = [[],[]]
    if orders:
        print "in_city"

    for order in orders:
        print "#############"
        print order.address
        #print get_pos2(order.address)
        postcode = order.address.split(',')[-1]
        if postcode:
            if in_pos(postcode,postcode_city,postcode_full_city):
                print "city:"
                print order
                return_list[index_city].append(order)
            else:
                print "other:"
                print order
                return_list[index_other].append(order)
        else:
            warning_window(master,"{} raises time out".format(order.address))
            return
       

    return return_list

# check postcode of various forms in list or not    
def in_pos(postcode,list1,list_full):
    num_form = []
    for i in postcode:
        if i.isdigit():
            num_form.append(i)
    pos_num = "".join(num_form)
    print "num_form:"
    print num_form
    
    # user enter number of postcode
    if num_form:
        if "".join(num_form) in list1:
            return True
        else:
            return False
    else:
        new_pos = rip_num(postcode)
        for pos_full in list_full:
            new_pos_full = rip_num_full(pos_full)
            print "new_pos_full:"
            print new_pos_full.lower()
            print new_pos.lower()
            if new_pos_full.lower() == new_pos.lower():
                print "$$$"
                return True
        return False
        
# get postcode from an address
def get_pos(address):
    try:
        print "gmaps.places_autocomplete(input_text = address, \
                                                 location = address)"
        print gmaps.places_autocomplete(input_text = address, \
                                                 location = address)
        print "gmaps.places_autocomplete(input_text = address, \
                                                 location = address)[0]"
        print gmaps.places_autocomplete(input_text = address, \
                                                 location = address)[0]
        new_location = gmaps.places_autocomplete(input_text = address, \
                                                 location = address)[0]['description']
        locationId = gmaps.geocode(new_location)[0]['place_id']
        fomattedAddress = gmaps.reverse_geocode(locationId)[0]['formatted_address']
        print fomattedAddress

        postcode = []
        for i in fomattedAddress.split(',')[-2]:
            if i.isdigit():
                postcode.append(i)
        print "get_pos"
        print postcode
        
        return "".join(postcode)
    except:
        print "Error Time Out"
        return

# get postcode from an address
def get_pos2(address):
    postcode = []
    for i in address.split(',')[1]:
        if i.isdigit():
            postcode.append(i)
    return "".join(postcode)

# compare postcode using name:
def get_pos_full(address):
    address = "VIC " + address +",Australia"
    print address
    new_location = gmaps.places_autocomplete(input_text = address, \
                                             location = address)[0]['description']
    print new_location
    return new_location
    
# once confirmed cakes are being processed,change the state of the order
# and update dispatcher info
def update_state2(num,orders,menu_var,window):
    for i in range(num):
        for order in orders[i]:
            # dispatcher_info = menu_var[i].get()
#             dispatcher_name = dispatcher_info.split(",")[index_dis_name]
#
#             for dis in db.dispatchers:
#                 if dis[index_dis_name] == dispatcher_name:
#                     try:
#                         dis[index_dis_orders].append(order.order_number)
#                     except AttributeError:
#                         list1 = convert_data_to_list(dis[index_dis_orders])
#                         list1.append(order.order_number)
#                         dis[index_dis_orders] = list1

            for item in db.orders_all:
                if order == item:
                    item.state = processing
                    if menu_var[i].get() != hint1:
                        item.set_dispatcher(menu_var[i].get())
                    else:
                        item.set_dispatcher("dispatcher not in database")

    db.update_dispatchers()
    db.update()
    refresh_all()
    window.destroy()

# a window contains warning messages
def warning_window(master,message):
    window = Toplevel(master)
    window.title("Warning")
    Label(window,text = message).pack()
    Button(window, text = "OK", command = window.destroy).pack()
    center(window)

#########################################################################################
#########################     cake   ####################################################
#########################################################################################
# add cake to the database
def add_cake(size,shape,inner,typee,cake_ps,window,list_box_cakes,cakes):
    if isvalid(cake_ps.get()):
        cake_no = (len(cakes)+1)
        cake_nos = [x[index_cake_no] for x in cakes]
        while cake_no in cake_nos:
            cake_no += 1
        new_cake = [cake_no,size.get(),shape.get(),inner.get(),typee.get(),cake_ps.get()]

        print "new_cake:"
        print new_cake
        # gobal list storing cakes info to be added or edited
        cakes.append(new_cake)
        print "current_cakes(add_cake)"
        print cakes
        
        list_box_cakes.insert_row(new_cake)
        last = len(list_box_cakes.table_data)-1
        tmp =  list_box_cakes.table_data[last]
        list_box_cakes.delete_row(last)
        list_box_cakes.insert_row(tmp)
        
        window.destroy()
        return True
    else:
        warning_window(master,"备注一栏不能含有除字母、汉字、数字、空格及逗号以外的字符！")
        return False
        
    
# add cake window
def window_add_cake(root,list_box_cakes,cakes):
    window = Toplevel(root)
    window.title("添加蛋糕")

    # set labels and entries(option menu)
    var_size,var_shape,var_inner,var_type,cake_ps,row_num = pack_cake_entry(window)
    
    print var_size.get()
    print var_shape.get()
    print var_inner.get()
    print var_type.get()
    print cake_ps.get()

    Button(window,text = "添加",command = lambda:add_cake(var_size,var_shape,\
            var_inner,var_type,cake_ps,window,list_box_cakes,cakes))\
            .grid(row=row_num,column = 0,pady = 4)
    Button(window,text= "取消",command = window.destroy).\
    grid(row=row_num,column = 1,pady = 4)
    center(window)
    

# pack text labels for text
def pack_cake_label(window):
    Label(window, text="尺寸:").grid(row=0,sticky=W,pady = 4)
    Label(window, text="形状:").grid(row=1,sticky=W,pady = 4)
    Label(window, text="内芯:").grid(row=2,sticky=W,pady = 4)
    Label(window, text="款式:").grid(row=3,sticky=W,pady = 4)
    Label(window, text="备注:").grid(row=4,sticky=W,pady = 4)

# set entries and labels for cake manipulation
def pack_cake_entry(window):
    pack_cake_label(window)
    
    #read size and type from file
    cake_size = read_cake_size()
    cake_shapes = read_cake_shapes()
    cake_types = read_cake_types()
    cake_inners = read_cake_inner()

    var_size = StringVar()
    var_size.set(cake_size[0])
    cake_size = OptionMenu(window,var_size, *cake_size)       # option menu for cake size
    cake_size.grid(row=0,column = 1, sticky=W,pady = 4)
    
    var_shape = StringVar()
    var_shape.set(cake_shapes[0])
    cake_shape = OptionMenu(window,var_shape,*cake_shapes)       # option menu for cake shape
    cake_shape.grid(row=1,column = 1, sticky=W,pady = 4)
    
    var_inner = StringVar()
    var_inner.set(cake_inners[0])
    cake_inner =OptionMenu(window,var_inner, *cake_inners)      # option menu for cake inners
    cake_inner.grid(row=2,column = 1, sticky=W,pady = 4)

    var_type = StringVar()
    var_type.set(cake_types[0])
    cake_type =OptionMenu(window,var_type, *cake_types)      # option menu for cake types
    cake_type.grid(row=3,column = 1, sticky=W,pady = 4)
    
    cake_ps = Entry(window)
    cake_ps.grid(row=4,column = 1, sticky=W,pady = 4)
    return var_size,var_shape,var_inner,var_type,cake_ps,5

# edit cake info
def edit_cake_window(master,list_box_cakes,cakes):
    if list_box_cakes.selected_rows:
        window = Toplevel(master)
        window.title("更改蛋糕信息")

        edit_cake = list_box_cakes.selected_rows[0]
        var_size,var_shape,var_inner,var_type,cake_ps,row_num = pack_cake_entry(window)

        var_size.set(edit_cake[index_cake_size])
        var_shape.set(edit_cake[index_cake_shape])
        var_inner.set(edit_cake[index_cake_inner])
        var_type.set(edit_cake[index_cake_type])
        cake_ps.insert(0,edit_cake[index_cake_ps])

        Button(window,text = "确认",command = lambda:edit_cake_info(\
            edit_cake,var_size,var_shape,var_inner,var_type,cake_ps,window,list_box_cakes,cakes))\
            .grid(row=row_num,column = 0,pady = 4)
        Button(window,text= "取消",command = window.destroy).\
            grid(row=row_num,column = 1,pady = 4)
        center(window)
          
    else:
        return

# edit cake info
def edit_cake_info(edit_cake,size,shape,inner,typee,
                              cake_ps,window,list_box_cakes,cakes):
    index = list_box_cakes.indices_of_selected_rows[0]
    new_cake = [edit_cake[index_cake_no],size.get(),shape.get(),inner.get(),typee.get(),cake_ps.get()]
    print "before:"
    print cakes
    new_cakes = []
    for cake in cakes:
        if edit_cake[index_cake_no] == cake[index_cake_no]:
            new_cakes.append(new_cake)
        else:
            new_cakes.append(cake)
    cakes = new_cakes
    list_box_cakes.update_row(index,new_cake)
    print "after:"
    print cakes
    window.destroy()
    
# window for deleting selected cake
def delete_cake_window(master,list_box_cakes,cakes):
    # print "\n\n\n########################\nlist_box_cakes.selected_rows"
 #    print list_box_cakes.selected_rows
 #    print list_box_cakes.selected_rows[0][0]
 #    print list_box_cakes.selected_rows[0][1]
 #    print list_box_cakes.selected_rows[0][2]
 #    print list_box_cakes.selected_rows[0][3]

    if list_box_cakes.selected_rows:
        edit_cake = list_box_cakes.selected_rows[0]

        window = Toplevel(master)
        window.config(width = 300,height = 80)
        center(window)
        window.title("删除蛋糕")
        Label(window, text="       确定删除这些蛋糕吗？       ").pack()
        Button(window, width = 8,text='确认', command= lambda:delete_cake(\
            window,edit_cake,list_box_cakes,cakes)).pack(pady=2)

        Button(window, width = 8,text='取消', command= \
            window.destroy).pack(pady=2)
    else:
        return

# delete selected cakes
def delete_cake(master,edit_cake,list_box_cakes,cakes):
    # print "edit cake :"
 #    print edit_cake
 #    print "######################"
 #    print "cakes"
 #    print cakes
    for data in list_box_cakes.selected_rows:
        for cake in cakes:
            if data[index_cake_no] == cake[index_cake_no]:
                cakes.remove(cake)
        
    # print "######################"
#     print "after:"
#     print cakes
    update_listbox(cakes,list_box_cakes)
    master.destroy()



# build a arranged list box according to row number
def build_list_box_cake(window,row_num):
    list_box_cakes = Multicolumn_Listbox(window, header_cake, \
                stripped_rows = ("white","#f2f2f2"), cell_anchor="center",height = 5)
                
    list_box_cakes.configure_column(0,width = 45)
    list_box_cakes.configure_column(1,width = 45)
    list_box_cakes.configure_column(2,width = 45)
    list_box_cakes.configure_column(3,width = 45)
    list_box_cakes.configure_column(4,width = 70)
    list_box_cakes.configure_column(5,width = 90)
    
    list_box_cakes.interior.grid(row=row_num, column=0,columnspan =2,rowspan = 3)
    list_box_cakes.interior.bind("<Button-2>", do_popup_cake)
    return list_box_cakes,(row_num+3)

#########################################################################################
#########################     order   ###################################################
#########################################################################################
# get address
def get_address(na_ad):
    tmp = na_ad.split(',')
    tmp.remove(na_ad.split(',')[0])
    #print tmp
    return tmp

# calculate the route and provide basic info
def run_order():
    if not db.orders_rightbox :
        warning_window(master,warning_message2)
        return
    elif not (proc_entry_num.get() or proc_entry_num_city.get()) :
        warning_window(master,warning_message5)
        return
    print "db.orders_rightbox"
    print db.orders_rightbox
    tmp = db.orders_rightbox
    print "tmp"
    print tmp 
    order_numbers = [">       选择订单        "]
    
    if proc_entry_num_city.get() and proc_entry_num.get():
        orders_areas = in_city(db.orders_rightbox,db.postcode[index_city],db.postcode_full[index_city])
        if orders_areas:
            random_order("Melbourne City",int(proc_entry_num_city.get()),orders_areas[index_city],order_numbers)
            random_order("South-east area",int(proc_entry_num.get()),orders_areas[index_other],order_numbers)
        else:
            warning_window(master,"Usage is up to today's limit")
            
    elif proc_entry_num_city.get() and (not proc_entry_num.get() or int(proc_entry_num.get()) == 0):
        random_order("Melbourne City",int(proc_entry_num_city.get()),tmp,order_numbers)
    elif proc_entry_num.get() and (not proc_entry_num_city.get() or int(proc_entry_num_city.get()) == 0):
        random_order("South-east area",int(proc_entry_num.get()),tmp,order_numbers)
  
#assign order randomly
def random_order(title,num_dis,orders_area,order_numbers):
    if orders_area:
        if not db.orders_rightbox:
            print "db.orders_rightbox null (random_order)"
        else:
            print "not null(random_order)"
        tmp = []
        for order in orders_area:
            tmp.append(order)
        random.seed(datetime.now())
        orders = []
        for i in range(num_dis):
            orders.append([])
        if not db.orders_rightbox:
            print "db.orders_rightbox null (random_order)"
        else:
            print "not null(random_order)"
        i = 0
        while tmp:
            if i >= num_dis:
                i = 0
            order = random.choice(tmp)
            orders[i].append(order)
            tmp.remove(order)
            i += 1
        if not db.orders_rightbox:
            print "db.orders_rightbox null (random_order)"
        else:
            print "not null(random_order)"
        manual_order(title,num_dis,tmp,orders,order_numbers)
    else:
        warning_window(master,title + warning_message7)

# allow user to manually assign orders
def manual_order(title,num_dis,orders_area,orders,order_numbers):
    manual_window = Toplevel(master)
    manual_window.title(title)

    if not db.orders_rightbox:
        print "db.orders_rightbox null (manual_order)"
    listboxes = []
    dispatchers = to_na_ad(db.dispatchers)
    dispatchers.insert(0,hint1)
    print "dispatchers"
    print dispatchers
    menu_var = []
    for i in range(num_dis):
        if i < 6:
            row_num = 0
            col_num = i
        elif 6<= i < 12:
            row_num = 4
            col_num = i-6
        elif 12 <= i < 18:
            row_num = 8
            col_num = i-12

        #set up dispatchers list menu
        var_dis = StringVar()
        var_dis.set(dispatchers[0])
        om_dispatchers = OptionMenu(manual_window,var_dis, *dispatchers)
        menu_var.append(var_dis)

        Label(manual_window,text = "派送员 No.{}".format((i+1))).grid(row=row_num,column=col_num,sticky = W)
        om_dispatchers.grid(row=(row_num+1),column=col_num,sticky = W)
        listbox = Multicolumn_Listbox(manual_window,
                            ['No'],
                            stripped_rows = ("white","#f2f2f2"),
                            cell_anchor="center",
                            height=5,
                            select_mode = EXTENDED)
        listbox.update(to_order2(orders[i]))
        justify_order2(listbox)
        listbox.interior.grid(row = (row_num+2),column =col_num)
        listboxes.append(listbox)
        Button(manual_window,text ="添加",width = 8,command =
                      lambda i=i,orders=orders,listboxes=listboxes,order_numbers=order_numbers
                      : manual_add_window(i,orders,listboxes,order_numbers,manual_window)).grid(row=(row_num+3),column=col_num,sticky = W)
        Button(manual_window,text ="删除",width = 8,command =
                      lambda i=i,orders=orders,listboxes=listboxes,order_numbers=order_numbers
                      : manual_delete(i,orders,listboxes,order_numbers)).grid(row=(row_num+3),column=col_num,sticky = E,padx=5)
        Button(manual_window,text ="打印文件并查看路线",command =
                      lambda i=i:cal_dis(orders[i],menu_var[i].get())).grid(row=(row_num+4),column=col_num,sticky = W)
    Button(manual_window,text ="确认所有订单",width = 12,
                 command = lambda i=num_dis: update_state2(i,orders,menu_var,manual_window)).grid(row=row_num+5,columnspan = (col_num+1))
    center(manual_window)

 # delete order in corresspoding listbox in the manul window
def manual_delete(i,orders,listboxes,order_numbers):
    if listboxes[i].selected_rows:
        selecteds = listboxes[i].selected_rows
        for selected in selecteds:
            print "i = {}".format(i)
            print selected
            order = get_item_by_num2(selected[0],db.orders_rightbox)
            orders[i].remove(order)
            order_numbers.append([order.order_number,order.address])
            print "manual_delete(order_numbers.append)"
            print order_numbers
            print listboxes[i].indices_of_selected_rows
            listboxes[i].delete_row(listboxes[i].indices_of_selected_rows[0])
            #listboxes[i].update(to_order2(orders[i]))
    else:
        return


# add order to corresspoding listbox in the manul window
def manual_add_window(i,orders,listboxes,order_numbers,window):
    manual_add_window = Toplevel(window)
    manual_add_window.title("添加订单")

    var = StringVar()
    var.set(order_numbers[0])
    om_order_numebrs = OptionMenu(manual_add_window,var, *order_numbers)

    om_order_numebrs.pack()
    Button(manual_add_window,width = 8,text ="确认",command =
                 lambda i=i,var=var,orders=orders,listboxes=
                 listboxes: manual_add(i,var,orders,listboxes,manual_add_window,order_numbers,om_order_numebrs)).pack()
    center(manual_add_window)

# add order to corresspoding listbox
def manual_add(i,var,orders,listboxes,manual_add_window,order_numbers,m):
    if var.get() != ">       选择订单        ":
        order = get_item_by_num(var.get(),db.orders_rightbox)
        orders[i].append(order)
        print "orders[i]"
        print order.address
        for k in orders[i]: 
            print k
        print order_numbers
        order_numbers.remove([order.order_number,order.address])

        menu = m.children['menu']
        # Clear the menu.
        menu.delete(0, 'end')
        for num in order_numbers:
            print num
            # Add menu items.
            menu.add_command(label=num, command=lambda: var.set(num))
        print "###########"
        print listboxes[i].table_data
        for j in orders[i]:
            print j.order_number
            
        listboxes[i].insert_row(to_order2([order])[0])
        last = len(listboxes[i].table_data)-1
        tmp =  listboxes[i].table_data[last]
        listboxes[i].delete_row(last)
        listboxes[i].insert_row(tmp)
        manual_add_window.destroy()


# get list item according to address and order num
def get_item(list,item):
    print "get_item"
    for item1 in list:
        if str(item1.order_number) == str(item[index_no]):
            print "get_item:"
            print item1
            return item1

# add left list box items to the right
def left_to_right():
    if not proc_left_box.indices_of_selected_rows:
        return

    selected_indices = proc_left_box.indices_of_selected_rows

    for index in selected_indices:
        listbox_item = proc_left_box.row_data(index)
        proc_right_box.insert_row(listbox_item)
        db.orders_rightbox.append(get_item(db.orders_leftbox,listbox_item))
        proc_left_box.delete_row(index)
        db.orders_leftbox.remove(get_item(db.orders_leftbox,listbox_item))

# add right list box items to the left
def right_to_left():
    if not proc_right_box.indices_of_selected_rows:
        return

    selected_indices = proc_right_box.indices_of_selected_rows

    for index in selected_indices:
        listbox_item = proc_right_box.row_data(index)
        proc_left_box.insert_row(listbox_item)
        db.orders_leftbox.append(get_item(db.orders_rightbox,listbox_item))
        proc_right_box.delete_row(index)
        db.orders_rightbox.remove(get_item(db.orders_rightbox,listbox_item))

# add all left list box items to the right
def add_all():
    for item in db.orders_leftbox:
        db.orders_rightbox.append(item)

    db.orders_leftbox[:] = []
    refresh_all()

# add order to the raw list
def add_order(order,agent,location,name,phone,candle,tableware,writing,price,
              var2,ps_info,window,dispatcher,current_cakes,
              mode,pickup_date,pickup_time,upfront,indicator = True):
    global window_open
    
    new_location = location.get()
    if isvalid_ad(location.get()) == 2:
        list1 = []
        for i in location.get():
            if i == '，':
                list1.append(',')
            else:
                list1.append(i)
        new_location = "".join(list1)       
         
    if (order.get() == '' or phone.get() == ''\
        or new_location == '' or current_cakes == []):
        warning_window(window,warning_message1)
        return False
    elif isvalid_ad(new_location) == 1:
        warning_window(window,warning_message8)
        return False
    elif isvalid_ad(new_location) < 0:
        warning_window(window,warning_message9)
        return False
    elif isvalid_id(order.get()):
        warning_window(window,"    订单号仅允许数字！    ")
        return False
    else:        
        if indicator:
            for order1 in db.orders_all:
                if order1.order_number == order.get():
                    warning_window(master,"     订单号已存在，请重新输入     ")
                    return False  
    
        new_order = Order(order.get(),agent.get(),new_location,name.get(),phone.get(),\
                           candle.get(),tableware.get(),writing.get(),price.get(),var2.get(),\
                           mode.get(),pickup_date.get(),pickup_time.get(),ps_info.get(),current_cakes,\
                           dispatcher,upfront.get())
        print new_order
        db.add(new_order)

        #update listbox after adding new order
        refresh_all()

        window.destroy()
        # clear this global list of cakes infro after adding/editing
        # action done
        window_open = False
        return True
        
# calendar window
def pop_calendar(window,pickup_date):
    cal_window = Toplevel(window)
    cal_window.title('Calendar')
    ttkcal = Calendar(master=window,firstweekday=calendar.SUNDAY)
    ttkcal.pack(expand=1, fill='both')
    Button(cal_window,text = 'Confirm',command = \
                          lambda:confirm_date(window,ttkcal,pickup_date)).pack()
# confirm canlendar
def confirm_date(window,ttkcal,pickup_date):
    window.destroy()
    pickup_date = ttkcal.selection()
    print pickup_date

# arrange all entries including text lables for windows
# main use is to avoid duplicate codes
def pack_order_entry(window,cakes,item,indicator):
    order = Entry(window)                  # text entry for order
    agent = Entry(window)                  # agent entry for order
    location = Entry(window)
    #location = AutocompleteEntry(window,indicator)    
    name = Entry(window)                   # name entry for order
    phone = Entry(window)                  # phone number
    candle = Entry(window)                 # candle info
    tableware = Entry(window)              # tableware info
    writing = Entry(window)                # writing info
    price = Entry(window)                  # price of order
    upfront = Entry(window)
    var2 = StringVar()                     # option menu for order state
    var2.set(states[0])
    state =OptionMenu(window,var2,*states)

    var_mode = StringVar()                 # option menu for mode
    var_mode.set(modes[0])
    mode =OptionMenu(window,var_mode,*modes)
    
    pickup_date = Entry(window)
    pickup_time = Entry(window)
    ps_info = Entry(window)                # plus info about the order
    
    if item:
        # put placeholders for the convenience of user to edit details of an order
        order.insert(0,item.order_number)
        agent.insert(0,item.agent)
        location.insert(0,item.address)
        name.insert(0,item.name)
        phone.insert(0,item.phone)
        candle.insert(0,item.candle)
        tableware.insert(0,item.tableware)
        writing.insert(0,item.writing)
        price.insert(0,item.price)
        upfront.insert(0,item.upfront)
        var2.set(item.state)
        ps_info.insert(0,item.ps_info)
        var_mode.set(item.mode)
        print "item.pickup_date"
        print item.pickup_date
        print "item.pickup_time"
        print item.pickup_time
        if item.pickup_date:
            pickup_date.insert(0,item.pickup_date)
        if item.pickup_time:
            pickup_time.insert(0,item.pickup_time)
    
    order.grid(row=0, column=1,sticky=W)
    agent.grid(row=1, column=1,sticky=W)
    location.grid(row=2, column=1,sticky=W)
    name.grid(row=3, column=1,sticky=W)
    phone.grid(row=4, column=1,sticky=W)
    candle.grid(row=5, column=1,sticky=W)
    tableware.grid(row=6, column=1,sticky=W)
    writing.grid(row=7, column=1,sticky=W)
    price.grid(row=8, column=1,sticky=W)
    upfront.grid(row=9, column=1,sticky=W)
    state.grid(row=10, column=1,sticky=W)
    mode.grid(row=11,column = 1,sticky=W)
    pickup_date.grid(row=12,column = 1,sticky=W)
    pickup_time.grid(row=13,sticky=W,column = 1)
    ps_info.grid(row=14,column = 1,sticky=W)
    list_box_cakes,row_num = pack_order_label(window)

    Button(window, text='添加蛋糕', command= \
            lambda:window_add_cake(window,\
            list_box_cakes,cakes),width = 8).grid(row=15, column=1,sticky=W)

    Button(window,text= "设置",width = 8,command = \
            lambda: cake_setting_window(window)).\
            grid(row=15,column = 1, sticky=E)

    row_num = (row_num)
    return order,agent,location,name,phone,candle,tableware,writing,price,var2,ps_info,var_mode,pickup_date,pickup_time,row_num,list_box_cakes,upfront

# pack all text labels
def pack_order_label(window):
    Label(window, text="订单号*:").grid(row=0,sticky=W)
    Label(window, text="客服号:").grid(row=1,sticky=W)  
    Label(window, text="地址(格式:St,postcode)*:").grid(row=2,sticky=W)
    Label(window, text="姓名:").grid(row=3,sticky=W)
    Label(window, text="电话*:").grid(row=4,sticky=W)
    Label(window, text="蜡烛:").grid(row=5,sticky=W)
    Label(window, text="餐具:").grid(row=6,sticky=W)
    Label(window, text="写字:").grid(row=7,sticky=W)    
    Label(window, text="价格:").grid(row=8,sticky=W)
    Label(window, text="定金:").grid(row=9,sticky=W)
    Label(window, text="订单状态*:").grid(row=10,sticky=W)
    Label(window, text="订单派送方式:").grid(row=11,sticky=W)
    Label(window, text="（如自取）自取日期:").grid(row=12,sticky=W)
    Label(window, text="（如自取）自取时间:").grid(row=13,sticky=W)
    Label(window, text="备注:").grid(row=14,sticky=W)
    Label(window, text="蛋糕信息*:").grid(row=15,sticky=W)
    list_box_cakes,row_num = build_list_box_cake(window,16)
    return list_box_cakes,(row_num+2)

# add order window
def window_add_order():
    global window_open,popup_cake
    if not window_open:
        #window_open = True
        window = Toplevel(master)
        window.title("添加订单")

        cakes = []
        # user inputs
        order,agent,location,name,phone,candle,tableware,\
        writing,price,var2,ps_info,mode,pickup_date,\
        pickup_time,row_num,list_box_cakes,upfront = pack_order_entry(window,cakes,None,True)

        # popup Menu for cake
        popup_cake = Menu(window,tearoff = 0)
        popup_cake.add_command(label="显示信息",command = \
            lambda:edit_cake_window(window,list_box_cakes,cakes))
        popup_cake.add_command(label="删除",command = \
            lambda:delete_cake_window(window,list_box_cakes,cakes))
        popup_cake.add_separator()

        Label(window, text = " （*）不得为空").\
            grid(row=row_num,columnspan = 2,pady = 4)
        Button(window, text="确认",width = 8,command= \
               lambda:add_order(order,agent,location,name,phone,\
                                candle,tableware,writing,price,\
                                var2,ps_info,window,None,cakes,\
                                mode,pickup_date,pickup_time,upfront)).\
               grid(row=(row_num+1), column=0,columnspan = 2, pady=4)
        center(window)

    else:
        return
    
# create a new window for editing an order
def edit_info(master,listbox):
    global window_open,popup_cake
    item = get_selected_order()
    if item and not window_open:
        #window_open = True
        check_select_state()
        window = Toplevel(master)
        window.title("订单详情")

        current_cakes = item.cake_type
        print "item.cake_type"
        print current_cakes

        # popup Menu for cake
        popup_cake = Menu(window,tearoff = 0)
        popup_cake.add_command(label="显示信息",command = \
            lambda:edit_cake_window(window,list_box_cakes,current_cakes))
        popup_cake.add_command(label="删除",command = \
            lambda:delete_cake_window(window,list_box_cakes,current_cakes))
        popup_cake.add_separator()

        # user inputs
        order,agent,location,name,phone,candle,tableware,\
        writing,price,var2,ps_info,mode,pickup_date,\
        pickup_time,row_num,list_box_cakes,upfront = pack_order_entry(window,current_cakes,item,False)

        update_listbox(sorted(current_cakes),list_box_cakes)
        Button(window, width = 8,text='确认', command= \
               lambda: edit_order(order,agent,location,name,phone,\
                                   candle,tableware,writing,price,
                                   var2,ps_info,item,window,current_cakes,\
                                   mode,pickup_date,pickup_time,item.dispatcher,\
                                   upfront,listbox,list_box_cakes)).\
                                   grid(row=row_num, columnspan=2, pady=4)
        center(window)
    else:
        return

# edit order info
def edit_order(order,agent,location,name,phone,
                candle,tableware,writing,price,
                var2,ps_info,item,window,current_cakes,
                mode,pickup_date,pickup_time,dis,upfront,
                listbox,list_box_cakes):
                
    new_location = location.get()
    if isvalid_ad(location.get()) == 2:
        list1 = []
        for i in location.get():
            if i == '，':
                list1.append(',')
            else:
                list1.append(i)
        new_location = "".join(list1)       
          
    if (order.get() == '' or phone.get() == ''\
        or new_location == '' or current_cakes == []):
        warning_window(window,warning_message1)
        return False
    elif isvalid_ad(new_location) == 1:
        warning_window(window,warning_message8)
        return False
    elif isvalid_ad(new_location) < 0:
        warning_window(window,warning_message9)
        return False
    elif isvalid_id(order.get()):
        warning_window(window,"    订单号仅允许数字！    ")
        return False
    
    indicator = True
    # order number didnt get changed
    if order.get() == item.order_number:
        indicator = False
        
    # if changed check whether this order number exists
    exist =False
    if indicator:
        for o in db.orders_all:
            if o.order_number == order.get():
                exist = True
                break
                       
    orders = []
    print "before:"
    for orde in db.orders_all:
        print orde
    # if exists stop the current process
    if exist:
        warning_window(master,"     订单号已存在，请重新输入     ")
        return
    else:
        for o in db.orders_all:
            if o.order_number == item.order_number:
                new_order = Order(order.get(),agent.get(),new_location,name.get(),phone.get(),\
                                   candle.get(),tableware.get(),writing.get(),price.get(),var2.get(),\
                                   mode.get(),pickup_date.get(),pickup_time.get(),ps_info.get(),\
                                   sorted(list_box_cakes.table_data),dis,upfront.get())
                orders.append(new_order)
            else:
                orders.append(o)
                
    db.orders_all = orders
    print "after:"
    for orde in db.orders_all:
        print orde
    db.update()
    refresh_all()
    window.destroy()
          
# multiple delete
def delete_selected(listbox):
    print "delete_selected"
    if listbox.indices_of_selected_rows:
        list_orders = []
        list_order_nums =[]
        selected_indices = listbox.indices_of_selected_rows
        for index in selected_indices:
            print listbox.row_data(index)
            order = get_item(db.orders_all,listbox.row_data(index))
            print order
            list_orders.append(order)
            if not (order.order_number in list_order_nums):
                list_order_nums.append(str(order.order_number))
        
        # notify user to make sure he/she confirm the deletion
        window = Toplevel(master)
        window.title("删除订单")
        Label(window, text="         删除订单号:" + \
              ",".join(list_order_nums)+"         \n"+\
              "     注意！删除订单会删除改列表内所有关于此订单的信息     ").pack()
        Button(window,width = 8, text='确定', command= lambda:delete_order3(list_orders,window)).pack(pady=2)
        Button(window,width = 8, text='取消', command= \
            window.destroy).pack(pady=2) 
        center(window)
        
        
    else:
        return

# delete an order
def delete_order(order,window):
    db.delete(order)
    refresh_all()
    window.destroy()

def delete_order2(order):
    db.delete(order)
    refresh_all()

# delete an order
def delete_order3(orders,window):
    for order in orders:
        db.delete(order)
    refresh_all()
    window.destroy()
    
# search window
def search_window():
    window = Toplevel(master)
    window.title("查找订单")
    
    entry = Entry(window)
    entry.grid(row=0,column = 0)
    Button(window,text ="查找",width = 12,command = lambda:search2(entry)).grid(row=0,column = 1,padx = 5)
    Button(window,text ="关闭",width = 8,command = window.destroy).grid(row=1,columnspan=2,pady=3)
    

# search an order by any attribute
def search2(entry):
    listbox = None
    deselect_all()
    inputt = entry.get()
    tabtitle = tabControl.tab(tabControl.select(), "text")
    print tabtitle
    print inputt
    if tabtitle == "处理订单":
        listbox = today_list_box_all
        orders = db.orders_today
    elif tabtitle == "历史订单":
        listbox = history_listbox
        orders = db.orders_history
    elif tabtitle == "配送订单":
        listbox = proc_left_box
        orders = db.orders_leftbox
    if inputt:
        index = 0
        for item in listbox.table_data:
            # search any characters might exist in the database
            pattern = re.compile('.*' + str(inputt).lower() + '.*')
            for sub_item in item:
                if re.match(pattern, str(sub_item)):
                    listbox.select_row(index)
            index += 1   
            
# deselect all selected items in processing tab
def deselect_proc():
    proc_left_box.deselect_all()
    proc_left_box.deselect_all()

# deselect all selected items in today tab
def deselect_today():
    today_list_box_all.deselect_all()
    today_list_box_processing.deselect_all()

# deselect all selected items in history tab
def deselect_his():
    history_listbox.deselect_all()

# find out which and where an item is selected (eg. index and listbox)
def find_selected_item():
    selected_indices = []
    return_list = []
    if today_list_box_all.indices_of_selected_rows:
       selected_indices = today_list_box_all.indices_of_selected_rows
       return_list = [selected_indices[0],today_list_box_all,db.orders_today]
       print "find_selected_item:"
       print  selected_indices[0]
    elif today_list_box_processing.indices_of_selected_rows:
       selected_indices = today_list_box_processing.indices_of_selected_rows
       return_list = [selected_indices[0],today_list_box_processing,db.orders_proc]
    elif history_listbox.indices_of_selected_rows:
      selected_indices = history_listbox.indices_of_selected_rows
      return_list = [selected_indices[0],history_listbox,db.orders_history]

    elif proc_right_box.indices_of_selected_rows:
        selected_indices = proc_right_box.indices_of_selected_rows
        return_list = [selected_indices[0],proc_right_box,db.orders_rightbox]
    elif proc_left_box.indices_of_selected_rows:
        selected_indices = proc_left_box.indices_of_selected_rows
        return_list = [selected_indices[0],proc_left_box,db.orders_leftbox]

    return return_list

# get info about selected item(order)
def get_selected_order():
    if find_selected_item():
        print "item found"
        selected_info = find_selected_item()
        listbox = selected_info[1]
        orders = selected_info[2]
        order = get_item(orders,listbox.row_data(selected_info[0]))
        print listbox.row_data(selected_info[0])
        return order
    else:
        return None

#refresh all listboxes
def refresh_all():
    deselect_all()
    update_dis_listbox()
    update_listbox(order2lists(db.orders_today),today_list_box_all)
    update_listbox(ctl_num_address(db.orders_leftbox),proc_left_box)
    update_listbox(ctl_num_address(db.orders_rightbox),proc_right_box)
    update_listbox(order2lists(db.orders_history),history_listbox)

def refresh_all2():
    deselect_all()
    db.update()
    update_dis_listbox()
    update_listbox(order2lists(db.orders_today),today_list_box_all)
    update_listbox(ctl_num_address(db.orders_leftbox),proc_left_box)
    update_listbox(ctl_num_address(db.orders_rightbox),proc_right_box)
    update_listbox(class_to_list_without_info(db.orders_history),history_listbox)

# deselect all selected items in all listboxes    
def deselect_all():
    proc_left_box_dis.deselect_all()
    today_list_box_processing.deselect_all()
    today_list_box_all.deselect_all()
    proc_left_box.deselect_all()
    proc_right_box.deselect_all()
    history_listbox.deselect_all()
    

#########################################################################################
#########################     menu   ####################################################
#########################################################################################
# bind popup menu to histotry listbox
def do_popup_his(event):
    # display the popup menu
    try:
        popup_his.tk_popup(event.x_root, event.y_root)
    finally:
        # make sure to release the grab (Tk 8.0a1 only)
        popup_his.grab_release()

# bind popup menu to today listbox
def do_popup(event):
    # display the popup menu
    try:
        popup.tk_popup(event.x_root, event.y_root)
    finally:
        # make sure to release the grab (Tk 8.0a1 only)
        popup.grab_release()
        
# bind popup menu to proc_left listbox
def do_popup_left(event):
    # display the popup menu
    try:
        popup_left.tk_popup(event.x_root, event.y_root)
    finally:
        # make sure to release the grab (Tk 8.0a1 only)
        popup_left.grab_release()

# bind popup menu to proc_rigtht listbox
def do_popup_right(event):
    # display the popup menu
    try:
        popup_right.tk_popup(event.x_root, event.y_root)
    finally:
        # make sure to release the grab (Tk 8.0a1 only)
        popup_right.grab_release()
    
# bind popup menu to cake listbox
def do_popup_cake(event):
    # display the popup menu
    try:
        popup_cake.tk_popup(event.x_root, event.y_root)
    finally:
        # make sure to release the grab (Tk 8.0a1 only)
        popup_cake.grab_release()

# bind popup menu to dis listbox
def do_popup_dis(event):
    # display the popup menu
    try:
        popup_dis.tk_popup(event.x_root, event.y_root)
    finally:
        # make sure to release the grab (Tk 8.0a1 only)
        popup_dis.grab_release()

# make sure only items of exact one listbox can be selected at one time
def check_select_state():
    tabtitle = tabControl.tab(tabControl.select(), "text")
    if tabtitle == "处理订单":
        deselect_proc()
        deselect_his()
    elif tabtitle == "历史订单":
        deselect_proc()
        deselect_today()
    elif tabtitle == "配送订单":
        deselect_his()
        deselect_today()
  

def do_deselect(event):
    print("Mouse position: (%s %s)" % (event.x, event.y))

#########################################################################################
#########################     dispatcher   ##############################################
#########################################################################################
# pack all text labels for dispatcher window
def pack_dis_label(window):
    Label(window, text="姓名*:").grid(row=0,sticky=W,pady = 4)
    Label(window, text="地址(St+postcode)*:").grid(row=1,sticky=W,pady = 4)
    Label(window, text="电话*::").grid(row=2,sticky=W,pady = 4)
    Label(window, text="备注:").grid(row=3,sticky=W,pady = 4)

# arrange all entries including text lables for windows
# main use is to avoid duplicate codes
def pack_dis_entry(window):
    name = Entry(window)                # anem entry for dispatcher
    home = Entry(window)                # location entry for dispatcher
    phone = Entry(window)               # phone entry for dispatcher
    ps_info = Entry(window)             # plus info about the dispatcher

    pack_dis_label(window)
    name.grid(row=0, column=1,sticky=W,pady = 4)
    home.grid(row=1, column=1,sticky=W,pady = 4)
    phone.grid(row=2, column=1,sticky=W,pady = 4)
    ps_info.grid(row=3, column=1,sticky=W,pady = 4)

    return name,home,phone,ps_info,4

# add dispatchers window
def add_dispatchers_window(master):
    window = Toplevel(master)
    window.title("添加配送员")
    window.config(width = 350,height = 185)
    center(window)

    name,home,phone,ps_info,row_num = pack_dis_entry(window)

    Button(window,text = "确认",width = 8,command = lambda:add_dispatcher(\
        name,home,phone,ps_info,window))\
        .grid(row=row_num,column = 0, sticky=W,pady = 4)
    Button(window,text= "取消",width = 8,command = window.destroy).\
        grid(row=row_num,column = 1, sticky=W,pady = 4)

# add dispatcher to the database
def add_dispatcher(name,home,phone,ps_info,window):
    new_dispatcher = [name.get(),home.get(),phone.get(),ps_info.get()]
    db.dispatchers.append(new_dispatcher)
    refresh_all()
    window.destroy()

# show dispatchers info on listbox
def update_dis_listbox():
    db.update_dispatchers()
    update_listbox(show_dis(db.dispatchers),proc_left_box_dis)
    
# edit dispatcher info
def edit_info_dis(master):
    if proc_left_box_dis.selected_rows:
        dis = proc_left_box_dis.selected_rows[0]
        print dis
        print dis[0]
        window = Toplevel(master)
        window.title("修改派送员信息")
        window.config(width = 350,height = 185)
        center(window)

        name,home,phone,ps_info,row_num = pack_dis_entry(window)
        name.insert(0,dis[index_dis_name])
        home.insert(0,dis[index_dis_home])
        phone.insert(0,dis[index_dis_phone])
        ps_info.insert(0,dis[index_dis_ps])

        Button(window,text = "确认",width = 8,command = lambda:edit_dispatcher(\
            name,home,phone,ps_info,window,dis))\
            .grid(row=row_num,column = 0, sticky=W,pady = 4)
        Button(window,text= "取消",width = 8,command = window.destroy).\
            grid(row=row_num,column = 1, sticky=W,pady = 4)
    else:
        return


# add dispatcher to the database
def edit_dispatcher(name,home,phone,ps_info,window,old_dis):
    old_dispatcher = find_dis(db.dispatchers,old_dis)
    db.dispatchers.remove(old_dispatcher)
    new_dispatcher = [name.get(),home.get(),phone.get(),ps_info.get()]
    db.dispatchers.append(new_dispatcher)
    refresh_all()
    window.destroy()

# delete dispatcher
def delete_window_dis():
    if proc_left_box_dis.selected_rows:
        selected_dis = proc_left_box_dis.selected_rows[0]
        disp = find_dis(db.dispatchers,selected_dis)
        db.dispatchers.remove(disp)
        db.update_dispatchers()
        update_listbox(db.dispatchers,proc_left_box_dis)        
    else:
        return

# only get first two data(name,address) to show in the listbox
def to_na_ad(old_list):
    dispatchers = []
    for dis in old_list:
        print "rip(dis[index_dis_name]"
        print rip(dis[index_dis_name])
        dispatchers.append("".join([rip(dis[index_dis_name]) + "," +dis[index_dis_home]]))
        # if dis[index_dis_home]:
        #dispatchers.append("".join([dis[index_dis_name] + "," +dis[index_dis_home]]))
    return dispatchers
#rip off comma
def rip(name):
    new_name = []
    for i in name:
        if not i.isalpha() and  not i.isdigit():
            new_name.append(' ')
        else:
            new_name.append(i)
    return "".join(new_name)

# find dispatcher
def find_dis(list,disp):
    for dis in list:
        if str(dis[index_dis_name]) == str(disp[index_dis_name])\
           and str(dis[index_dis_home]) == str(disp[index_dis_home]):
           return dis
    return None

# build x scrollbar and y scrollbar for each tab
def build_scrollbar(container):
    canvas = Canvas(container, highlightthickness=0)
    xscroll = Scrollbar(container,orient="horizontal",command=canvas.xview)
    yscroll = Scrollbar(container,command=canvas.yview)
    canvas.config(xscrollcommand=xscroll.set,
                  yscrollcommand=yscroll.set,
                  scrollregion=(0,0,1400,500),
                  width=1200,height=490)

    xscroll.pack(side=BOTTOM, fill=X)
    yscroll.pack(side=RIGHT, fill=Y)
    canvas.pack(side=LEFT, fill=BOTH, expand=True)
    tab = Frame(canvas)
    canvas.create_window((0,0),window=tab,anchor='nw')

    return tab

#########################################################################################
#########################     settings   ################################################
#########################################################################################
#  setting window to modify the option of cake size, cake type etc.
def cake_setting_window(root):
    window = Toplevel(root)
    window.title("设置")
    tmp_size = read_cake_size()
    tmp_type = read_cake_types()
    tmp_inner = read_cake_inner()
    tmp_shapes = read_cake_shapes()

    # listbox of size
    size_listbox = Multicolumn_Listbox(window,
                        ['尺寸'],
                        stripped_rows = ("white","#f2f2f2"),
                        cell_anchor="center",
                        height=5)

    size_listbox.update(to_listbox(tmp_size))
    size_listbox.interior.grid(row = 0,column =0)
    Button(window,text ="添加",width = 8,command = lambda: cake_setting_add_window(\
                        window,"尺寸",size_listbox,tmp_size)).grid(row = 1,column =0,sticky=W)
    Button(window,text ="删除",width = 8,command = lambda: cake_setting_delete(\
                        size_listbox,tmp_size)).\
                        grid(row = 1,column =0,sticky=E)

    # listbox of cake type
    type_listbox = Multicolumn_Listbox(window,
                        ['款式'],
                        stripped_rows = ("white","#f2f2f2"),
                        cell_anchor="center",
                        height=5)
    type_listbox.update(to_listbox(tmp_type))
    type_listbox.interior.grid(row = 0,column =1)
    Button(window,text ="添加",width = 8,command = lambda: cake_setting_add_window(\
                        window,"款式",type_listbox,tmp_type)).grid(row = 1,column =1,sticky=W)
    Button(window,text ="删除",width = 8,command = lambda: cake_setting_delete(\
                        type_listbox,tmp_type)).grid(row = 1,column =1,sticky=E)
                        
    # listbox of cake inner
    inner_listbox = Multicolumn_Listbox(window,
                        ['内芯'],
                        stripped_rows = ("white","#f2f2f2"),
                        cell_anchor="center",
                        height=5)
    inner_listbox.update(to_listbox(tmp_inner))
    inner_listbox.interior.grid(row = 0,column =2)
    Button(window,text ="添加",width = 8,command = lambda: cake_setting_add_window(\
                        window,"内芯",inner_listbox,tmp_inner)).grid(row = 1,column =2,sticky=W)
    Button(window,text ="删除",width = 8,command = lambda: cake_setting_delete(\
                        inner_listbox,tmp_inner)).grid(row = 1,column =2,sticky=E)
                        
    # listbox of cake inner
    shapes_listbox = Multicolumn_Listbox(window,
                        ['形状'],
                        stripped_rows = ("white","#f2f2f2"),
                        cell_anchor="center",
                        height=5)
    shapes_listbox.update(to_listbox(tmp_shapes))
    shapes_listbox.interior.grid(row = 0,column =3)
    Button(window,text ="添加",width = 8,command = lambda: cake_setting_add_window(\
                        window,"内芯",shapes_listbox,tmp_shapes)).grid(row = 1,column =3,sticky=W)
    Button(window,text ="删除",width = 8,command = lambda: cake_setting_delete(\
                        shapes_listbox,tmp_shapes)).grid(row = 1,column =3,sticky=E)
    Button(window,text ="确认",width = 8,command = \
                        lambda: cake_setting_confirm(window,tmp_size,tmp_type,tmp_inner,tmp_shapes)).grid(row = 2,columnspan =4,pady =4)
    
    center(window)

# confirm all the changes
def cake_setting_confirm(window,size,types,inner,shapes):
     f = open(file_size,'wb')
     f.truncate()
     writer = csv.writer(f,delimiter = ',')
     writer.writerow(size)
     f.close()

     f = open(file_type,'wb')
     f.truncate()
     writer = csv.writer(f,delimiter = ',')
     writer.writerow(types)
     f.close()

     f = open(file_inner,'wb')
     f.truncate()
     writer = csv.writer(f,delimiter = ',')
     writer.writerow(inner)
     f.close()
     
     f = open(file_shapes,'wb')
     f.truncate()
     writer = csv.writer(f,delimiter = ',')
     writer.writerow(shapes)
     f.close()
     
     window.destroy()
# add size/type
def cake_setting_add_window(window,indicator,listbox,list1):
    cake_setting_add_window = Toplevel(window)
    cake_setting_add_window.title("添加")
    Label(cake_setting_add_window,text = "新{}:".format(indicator)).\
                     grid(sticky =W,row=0,column =0,pady = 4)
    entry = Entry(cake_setting_add_window)
    entry.grid(sticky =W,row=0,column =1,pady = 4)
    Button(cake_setting_add_window,text = "确认",command = \
                     lambda :cake_setting_add(indicator,listbox,entry,cake_setting_add_window,list1)).\
                     grid(row=1,columnspan =2,pady = 4)
    center(cake_setting_add_window)

# add it to the database
def cake_setting_add(indicator,listbox,entry,window,list1):
    new_size = entry.get()
    print "cake_setting_add"
    print new_size
    list1.append(new_size)

    listbox.insert_row([new_size])
    last = len(listbox.table_data)-1
    tmp =  listbox.table_data[last]
    listbox.delete_row(last)
    listbox.insert_row(tmp) 
    window.destroy()

# delete size/type
def cake_setting_delete(listbox,list1):
    if listbox.selected_rows:
        print listbox.selected_rows
        print list1
        for selected in listbox.selected_rows:
            try:
                list1.remove(str(selected[0]))
            except UnicodeEncodeError:
                list1.remove(selected[0])
        print list1
        
        listbox.update(to_listbox(list1))
    else:
        return

# about window
def about():
    window = Toplevel(master)
    window.title("关于")
    image = Image.open(logo)
    image = image.resize((200, 200), Image.ANTIALIAS)
    photo = ImageTk.PhotoImage(image)
    labimg = Label(window, image=photo)
    labimg.image = photo
    labimg.pack(fill = "both", expand = "yes")
    Label(window,text ="Copyright 2017-2018 Happy Hackers Pty Ltd. All rights reserved.").pack()

# open hh website
def open_web():
    webbrowser.open("http://happyhackers.com.au")

# move all selected items to required state
def move2state(listbox,state):
    if listbox.selected_rows:
        for line in listbox.selected_rows:
            for order in db.orders_all:
                print int(line[0])
                print int(order.order_number)
                if int(line[0]) == int(order.order_number):
                    order.state = state
    db.update()
    refresh_all()
     
# buld table
def build_table(window):
    listbox = Multicolumn_Listbox(window, header_new, \
                            stripped_rows = ("white","#f2f2f2"), \
                            cell_anchor="center",height=24,select_mode = EXTENDED)
    listbox.configure_column(0,width = col_size1)
    listbox.configure_column(1,width = col_size1)
    listbox.configure_column(2,width = col_size1)
    listbox.configure_column(3,width = col_size1)
    listbox.configure_column(4,width = col_size1)
    listbox.configure_column(5,width = col_size1)
    listbox.configure_column(6,width = col_size3)
    listbox.configure_column(7,width = col_size1)
    listbox.configure_column(8,width = col_size1)
    listbox.configure_column(9,width = col_size2)
    listbox.configure_column(10,width = col_size1)
    listbox.configure_column(11,width = col_size2)
    listbox.configure_column(12,width = col_size2)
    listbox.configure_column(13,width = (col_size2+250))
    listbox.configure_column(14,width = col_size2)
    listbox.configure_column(15,width = col_size3)
    return listbox
    
# buld table of history    
def build_table_his():
    listbox = Multicolumn_Listbox(tab_history,header_full, \
                                stripped_rows = ("white","#f2f2f2"), cell_anchor="center"\
                                              ,height=24,select_mode = EXTENDED)
    size = col_size3 +30
    listbox.configure_column(0,width = size)
    listbox.configure_column(1,width = size)
    listbox.configure_column(2,width = size)
    listbox.configure_column(3,width = size)
    listbox.configure_column(4,width = size)
    listbox.configure_column(5,width = size)
    listbox.configure_column(6,width = size)
    listbox.configure_column(7,width = size)
    listbox.configure_column(8,width = (size+50))
    listbox.configure_column(9,width = (size+50))
    return listbox
# check location validity
def check_location():
    locations = [x.address for x in db.orders_leftbox]
    # locs = [x.address for x in db.orders_rightbox]
#     locations.append(locs)
    isgood = 0
    
    if proc_start_loc.get():
        des = gmaps.places_autocomplete(input_text = proc_start_loc.get())
        if not des:
            warning_window(master," 初始地址可能存在问题 " +"\n请手动检查地址信息的正确性")
            isgood = -1
        
    for loc in locations:
        print loc
        des = gmaps.places_autocomplete(input_text = loc)
        if not des:
            warning_window(master," 这个地址可能存在问题 " + loc+"\n请手动检查地址信息的正确性")
            isgood = -1
        else:
          new_loc = gmaps.places_autocomplete(input_text = loc)[0]['description']
          if not new_loc:
              warning_window(master," 这个地址可能存在问题 " + loc+"\n请手动检查地址信息的正确性")
              isgood = -1
          else:
              new_loc_id = gmaps.geocode(new_loc)[0]['place_id']
              formatted_address = gmaps.reverse_geocode(new_loc_id)[0]['formatted_address']
              loc_id = gmaps.geocode(formatted_address)[0]['place_id']
              #print formatted_address + " " + loc_id
              #print new_loc + " " + new_loc_id
              if loc_id != new_loc_id:
                  warning_window(master," 这个地址可能存在问题 " + loc +"\n请手动检查地址信息的正确性")
                  isgood = -1
    if isgood == 0:            
        warning_window(master,"     暂未发现地址信息内存在任何问题     ")
    
# autocomplete entry 
# @source: http://code.activestate.com/recipes/578253-an-entry-with-autocompletion-for-the-tkinter-gui/
# modified by Rondo
class AutocompleteEntry(Entry):
    def __init__(self,window,indicator):
        
        Entry.__init__(self,window)
        self.window = window
        self.indicator = indicator
        self.var = self["textvariable"]
        if self.var == '':
            self.var = self["textvariable"] = StringVar()

        self.var.trace('w', self.changed)
        self.bind("<Right>", self.selection)
        self.bind("<Up>", self.up)
        self.bind("<Down>", self.down)
        
        self.lb_up = False

    def changed(self, name, index, mode):
        
        if not self.indicator:
            self.indicator = True
            return  
            
        if self.var.get() == '' or len(self.var.get()) < 8:
            try:
                self.lb.destroy()
            except:
                print "AttributeError"
            self.lb_up = False
        else:
            words = self.get_loc(self.var.get())
            if words:            
                if not self.lb_up:
                    self.lb = Listbox(self.window)
                    self.lb.config(width=0)
                    self.lb.bind("<Double-Button-1>", self.selection)
                    self.lb.bind("<Return>", self.selection)
                    self.lb.place(x=(self.winfo_x()-150), y=self.winfo_y()+self.winfo_height())
                    self.lb_up = True
                
                self.lb.delete(0, END)
                for w in words:
                    self.lb.insert(END,w)
            else:
                if self.lb_up:
                    self.lb.destroy()
                    self.lb_up = False
        
    def selection(self, event):

        if self.lb_up:
            self.var.set(self.lb.get(ACTIVE))
            self.lb.destroy()
            self.lb_up = False
            self.icursor(END)

    def up(self, event):

        if self.lb_up:
            if self.lb.curselection() == ():
                index = '0'
            else:
                index = self.lb.curselection()[0]
            if index != '0':                
                self.lb.selection_clear(first=index)
                index = str(int(index)-1)                
                self.lb.selection_set(first=index)
                self.lb.activate(index) 

    def down(self, event):

        if self.lb_up:
            if self.lb.curselection() == ():
                index = '0'
            else:
                index = self.lb.curselection()[0]
            if index != END:                        
                self.lb.selection_clear(first=index)
                index = str(int(index)+1)        
                self.lb.selection_set(first=index)
                self.lb.activate(index) 
                
    def get_loc(self,txt):
        new_loc = []
        for loc in gmaps.places_autocomplete(input_text = txt):
            print loc['description']
            tmp = loc['description'].split(",")
            print tmp
            new_loc.append(tmp[0] + "," + tmp[1])

        #new_loc = loc.remove(loc.split(',')[-1])
        return new_loc
        
############## MAIN FUNCTION ##################################################
############## MAIN FUNCTION ##################################################
############## MAIN FUNCTION ##################################################
############## MAIN FUNCTION ##################################################
if __name__ == "__main__":  
    # url = shop_url + "/orders.json"
    # print url
    # r = requests.get(url)
    # orders = r.json()['orders']
    # for order in orders:
    #     print "id = " + str(order['id']) + " " + str(order['fulfillment_status'])
    #     for line in order['line_items']:
    #         print "  "+line['title']+" "+line['variant_title']
 
##############################################################################    
    # main window setup
    master = Tk()
    master.title("iCake")
    
    # scrollbar = Scrollbar(master)
#     scrollbar.pack(side=BOTTOM, fill=X)

    db = []                              # orders database
    locations = []                       # list of locations from googlemaps api
    orders = {}                          # dictionary of ({id:info})
    infos = []                           # info of an order(ps_info)
    start_locationId = None              # the initial start location
    tabControl = ttk.Notebook(master)    # tab controller
    dis_right = []                       # list of dispatchers in the rilistbox
    tab_today = None
    tab_history = None
    tab_proc = None
    tab_dis = None

    db = Orders_database(file_order)
    #tab for 'Today','History','Processing'
    tab_today_container =Frame(tabControl,relief=GROOVE,width=50,height=100,bd=0)
    tab_today_container.place(x=10,y=10)
    tabControl.add(tab_today_container, text='处理订单')

    tab_history_container =Frame(tabControl,relief=GROOVE,width=50,height=100,bd=0)
    tab_history_container.place(x=10,y=10)
    tabControl.add(tab_history_container, text='历史订单')

    tab_proc_container =Frame(tabControl,relief=GROOVE,width=50,height=100,bd=0)
    tab_proc_container.place(x=10,y=10)
    tabControl.add(tab_proc_container, text='派送订单')

    tab_dis_container =Frame(tabControl,relief=GROOVE,width=50,height=100,bd=0)
    tab_dis_container.place(x=10,y=10)
    tabControl.add(tab_dis_container, text ='配送员')

    # build x scrollbar and y scrollbar for each tab
    # if statement to speed up the loading when switching tabs
    if not tab_today:
        tab_today = build_scrollbar(tab_today_container)
    if not tab_history:
        tab_history = build_scrollbar(tab_history_container)
    if not tab_dis:
        tab_proc = build_scrollbar(tab_proc_container)
    if not tab_dis:
        tab_dis = build_scrollbar(tab_dis_container)

    # background_image= PhotoImage("1.png")
    # background_label = Label(tab_today, image=background_image)
    # background_label.place(x=0, y=0, relwidth=1, relheight=1)

    # popup Menu
    popup = Menu(master,tearoff = 0)
    popup.add_command(label="显示信息",command = lambda:edit_info(master,today_list_box_all))
    popup.add_command(label="删除",command = lambda:delete_selected(today_list_box_all))
    # popup.add_command(label="移到历史",command = lambda:move2state(today_list_box_all,arrived))
    # popup.add_command(label="改成（等待被派送）",command = lambda:move2state(today_list_box_all,waiting))
    # popup.add_command(label="改成（正在派送）",command = lambda:move2state(today_list_box_all,processing))
    popup.add_command(label="刷新",command = lambda:refresh_all())
    
    popup.add_separator()
    
    # popup Menu for left listbox of processing tab
    popup_left = Menu(master,tearoff = 0)
    popup_left.add_command(label="显示信息",command = lambda:edit_info(master,today_list_box_all))
    popup_left.add_command(label="删除",command = lambda:delete_selected(proc_left_box))
    #popup_left.add_command(label="移到历史",command = lambda:move2state(proc_left_box,arrived))
    # popup_left.add_command(label="改成（等待被派送）",command = lambda:move2state(proc_left_box,waiting))
    # popup_left.add_command(label="改成（正在派送）",command = lambda:move2state(proc_left_box,processing))
    popup_left.add_command(label="刷新",command = lambda:refresh_all2())
    popup_left.add_separator()
    
    # popup Menu for left listbox of processing tab
    popup_right = Menu(master,tearoff = 0)
    popup_right.add_command(label="刷新",command = lambda:refresh_all2())
    popup_right.add_separator()

    # popup Menu for dispatcher
    popup_dis = Menu(tab_proc,tearoff = 0)
    popup_dis.add_command(label="显示信息",command = lambda:edit_info_dis(master))
    popup_dis.add_command(label="删除",command = lambda:delete_window_dis())
    popup_dis.add_command(label="刷新",command = lambda:refresh_all())
    popup_dis.add_separator()

    popup_his = Menu(master,tearoff = 0)
    popup_his.add_command(label="显示信息",command = lambda:edit_info(master,history_listbox))
    popup_his.add_command(label="删除",command = lambda:delete_selected(history_listbox))
    # popup_his.add_command(label="改成（等待被派送）",command = lambda:move2state(history_listbox,waiting))
    # popup_his.add_command(label="改成（正在派送）",command = lambda:move2state(history_listbox,processing))
    popup_his.add_command(label="刷新",command = lambda:refresh_all())
    popup_his.add_separator()

    # menu for app about
    menubar = Menu(master)
    # create a pulldown menu, and add it to the menu bar
    filemenu = Menu(menubar, tearoff=0)
    filemenu.add_command(label="关于",command=about)
    filemenu.add_command(label= "我们的官网",command=open_web)
    filemenu.add_separator()
    filemenu.add_command(label="退出", command=master.quit)
    menubar.add_cascade(label="App", menu=filemenu)
    
    searchmenu = Menu(menubar, tearoff=0)
    searchmenu.add_command(label="查找",command=search_window)
    menubar.add_cascade(label="Search", menu=searchmenu)
    


    master.config(menu=menubar)
#########################################TAB1############################################
    # set up ui for 'Today'
    today_search = Entry(tab_today)
    today_add_order = Button(tab_today, text = '添加订单',command = window_add_order,width = 15)
    today_button_refresh = Button(tab_today, text = '刷新',command =\
                                  refresh_all,width = 15)
    today_list_box_all = build_table(tab_today)
    
    today_list_box_processing = Multicolumn_Listbox(\
                            tab_today, header_proc, stripped_rows = \
                            ("white","#f2f2f2"), cell_anchor="center",height=7)
    today_text_orders_all = Label(tab_today, text="正在处理:",font=("Calibri",title_size))
    today_text_processing = Label(tab_today, text="正在派送:",font=("Calibri",title_size))


    # bind listbox to popup menu
    today_list_box_all.interior.bind("<Button-2>", do_popup)
    today_list_box_all.interior.bind("<Button-1>", do_deselect)


    today_text_orders_all.grid(row=0,column=0,sticky = W,padx =4)
    #today_search.grid(row=0,column=4, sticky = E)
    #today_button_search.grid(row = 0,column =5, sticky = W)
    #today_button_multiple_delete.grid(row=0,column=0, sticky = E)
    today_add_order.grid(row=0,column=7,sticky = E)
    #today_text_processing.grid(row=2,sticky = W,padx = 4)
    today_list_box_all.interior.grid(row=1,pady=4,column = 0,columnspan = 10)
    # today_list_box_processing.interior.grid(row=3,pady=4,column = 0,columnspan = 10)

########################################TAB2#############################################
    # set up ui for 'History'
    # history_date = Entry(tab_history)
    history_search = Entry(tab_history)

    # history_button_search = Button(tab_history, text = '查找',width = 15,command =\
    #                              lambda:search_listbox_history(history_search.get()))
    # history_listbox = Listbox(tab_history, width = 70,height = 20)
    history_listbox = build_table(tab_history)
    history_listbox.interior.bind("<Button-2>", do_popup_his)
    #history_date.grid(row = 0,column = 0,sticky = W)
    Label(tab_history,text = '历史订单',font=("Calibri",title_size)).\
                      grid(row=0,column = 0,sticky = W,padx=4)
    #history_search.grid(row=0,column = 4,sticky = E,padx=4)
    #history_button_search.grid(row = 0,column = 5,sticky=W)
    history_listbox.interior.grid(row=1,column = 0,columnspan = 10,pady=4)

#########################################TAB3############################################
    # set up ui for 'Processing'
    #orders ui elements
    proc_text_name = Label(tab_proc, text ="派送订单",font = ("Calibri",title_size))
    proc_entry_num = Entry(tab_proc)
    proc_entry_num_city = Entry(tab_proc)
    proc_entry_num.insert(0,1)
    proc_left_box = Multicolumn_Listbox(tab_proc, header_small, \
                        stripped_rows = ("white","#f2f2f2"), cell_anchor="center",\
                                        height=22,select_mode = EXTENDED)
    proc_right_box = Multicolumn_Listbox(tab_proc, header_small, \
                        stripped_rows = ("white","#f2f2f2"), cell_anchor="center",\
                                        height=22)
    proc_add_all_button = Button(tab_proc, text = "全选",command = add_all,width=4,padx=2)
    proc_add_button = Button(tab_proc, text = ">",command = left_to_right,width=4,padx=2)
    proc_delete_button = Button(tab_proc, text = "<",command = right_to_left,width=4,padx=2)

    proc_run_button = Button(tab_proc, text = "运行",command = run_order,\
                             width=15)
    proc_check_button = Button(tab_proc, text = "检查地址",command = check_location,\
                             width=15)
    #Label(tab_proc,text = "（仅等待被派送列表内）").grid(row = 9,column =5,sticky = W,padx = 30)
    # proc_manually_button = Button(tab_proc, text = "manually assgin orders",command = manually_run,\
    #                          width=20)
    proc_start_loc = Entry(tab_proc)
    proc_end_loc = Entry(tab_proc)
    #proc_start_loc.insert(0,start_location)

    #orders ui elements setup
    proc_left_box.interior.bind("<Button-2>", do_popup_left)
    proc_right_box.interior.bind("<Button-2>", do_popup_right)
    tab_proc.bind("<Button-2>", do_popup_right)
    
    proc_text_name.grid(row = 0, column = 0,sticky=W)
    Label(tab_proc, text ="等待被派送:").grid(row=1,column = 0,pady=4,sticky = W)
    proc_left_box.interior.grid(row=2,rowspan = 10,column = 0,columnspan =2)
    Label(tab_proc, text ="准备被派送:").grid(row=1,column = 3,pady=4,sticky = W)
    proc_right_box.interior.grid(row=2,rowspan = 10,column = 3,columnspan =2)
    proc_add_all_button.grid(row =5, column =2)
    proc_add_button.grid(row =4,column = 2)
    proc_delete_button.grid(row =6,column = 2)

    Label(tab_proc, text = "起始地址:").grid(row = 1,column =5,padx = 30,sticky = W)
    proc_start_loc.grid(row=2,column = 5,sticky = W,padx = 30,ipady = 10)
    Label(tab_proc, text = "东南区派送员数量:").grid(row = 3,column =5,padx = 30,sticky = W)
    proc_entry_num.grid(row=4,column = 5,sticky = W,padx = 30)
    Label(tab_proc, text = "City派送员数量:").grid(row = 5,column =5,padx = 30,sticky = W)
    proc_entry_num_city.grid(row=6,column = 5,sticky = W,padx = 30)

    proc_run_button.grid(row = 7,column =5,sticky = W,padx = 30)
    proc_check_button.grid(row = 8,column =5,sticky = W,padx = 30)
    # Label(tab_proc, text = "Or").grid(row = 8,column =5,padx = 30,sticky = W)
 #    proc_manually_button.grid(row = 9,column =5,padx = 30,sticky = W)

#########################################TAB4############################################
    # set up ui for dispatcher
    #dispatcher
    proc_left_box_dis = Multicolumn_Listbox(tab_dis, header_dispatcher_full, \
                        stripped_rows = ("white","#f2f2f2"), cell_anchor="center",\
                                        height=22)
    proc_left_box_dis.interior.bind("<Button-2>", do_popup_dis)
    proc_add_dispatcher_button = Button(tab_dis, text = "添加新派送员",\
                        command = lambda: add_dispatchers_window(tab_proc),width = 15)

    Label(tab_dis, text ="管理派送员",font=("Calibri",title_size))\
                .grid(row=0,column = 0,pady=4,sticky = W)
    proc_left_box_dis.interior.grid(row=1,rowspan = 7,column = 0,columnspan =5)
    proc_add_dispatcher_button.grid(row =0,column=4,sticky = E,pady = 4)

    #image
    image2 = Image.open(logo)
    image2 = image2.resize((200, 200), Image.ANTIALIAS)
    photo2 = ImageTk.PhotoImage(image2)
    labimg2 = Label(tab_dis, image=photo2)
    labimg2.image = photo2
    #image

    labimg2.grid(row = 3,column =5,padx = 30,sticky = S)
    Label(tab_dis, text = "如有任何问题，请发送邮件至 \nadmin@happyhackers.com.au",font=("Calibri",17)).grid(row = 6,column =5,padx = 30,sticky = S)
    Label(tab_dis, text = "This application is designed by  Happy Hackers Pty Ltd. \nAll rights reserved",font=("Calibri",15),foreground = 'gray').grid(row = 7,column =5,padx = 30,sticky = N)
    refresh_all()

    tabControl.pack(fill=BOTH, expand=YES)
    mainloop( )




